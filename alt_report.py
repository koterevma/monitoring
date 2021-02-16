from numpy.core.fromnumeric import size
import pandas as pd
import datetime
import docx.enum.text
import urllib.request
import odf
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def main():

    #INITIALIZATION PART 0

    brigades = ['Цыбулин П. А.\nМарков Г. С.\n', 'Базылев А. А.\nЯшин Н. А.\n', 'Патокин С. В.\nЗаидов А. Р.\n', 'Савченко В. А.\nЗайцев А. И.\n',
                    'Ольховик Н. С.\nЩербаков В. Ю.\n','Носков М. П.\nЦветков Ю. А.\n', 'Чупис С. А.\nФролова П. В.\n', 'Носков И. В.\nЛегин Д. И.\n',
                    'Красовский Г. О.\nШнырина А. Н.\n', 'Тебеньков К. А.\nМашкарь М. Ю.\n', 'Воскресенский О. А.\nАртёмов А. А.\n', 'Котерев М. А.\nМихальченко М. Е.\n',
                    'Кетрой Д. А.\nЛинкевич М. М.\n', 'Соловьев П. А.\nЗавьялов И. А.\n', 'Лямин А. С.\nСанжурин С. В.\n', 'Фалин А. Д.\nПолторацкий С. А.\n']
    working_directory = Path(__file__).parent

    today_date = datetime.date.today().strftime("%d-%m-%y-")
    path_file_in = working_directory / f"data/{today_date}info.csv"
    path_file_out = working_directory / f"reports/{today_date}report.docx"
    url = 'https://vk.com/doc318247198_588372326?hash=2935512c2301060857&dl=1bd55339548be8720b'
    #url = 'https://drive.google.com/u/0/uc?id=1dgdUBBmExKNNZmom_JVUfgHPXcwTN9bk&export=download' old link
    currbrigade = list()

    #TABLE READ PART 1

    with open(path_file_in, 'r', encoding="utf-8") as file:
        df = pd.read_csv(file, delimiter=';')

    df = df.set_index('Date')
    df.index = pd.to_datetime(df.index)

    curr_row = df.iloc[0]
    dates = [ df.index[0] ]
    values = [ list(curr_row) ]

    for idx, row in df.iterrows():
        if list(curr_row) != list(row):
            dates.append(idx)
            curr_row = row
            values.append(list(row))
    dates.append(df.index[len(df.index) - 1])

    doc = Document()
    tables = list()
    shade_row = list()
    shade_col = list()

    #FILE DOWNLOAD AND BRIGADE READ PART 2

    to_downloaded = working_directory / "temp/sch.ods"
    urllib.request.urlretrieve(url, to_downloaded)

    schtab = pd.read_excel(to_downloaded, 'График', engine="odf")

    clm = 0
    current_date = str(dates[0])[0:10]+' 00:00:00'

    for i in range(schtab.shape[1]):
        if(str(schtab.iloc[1,i])==current_date):
            clm = i
            break

    if clm > 0:
        thisday = list(schtab.iloc[:,clm])
        for i in range(len(thisday)):
            if thisday[i] >= 0 and thisday[i] <= 5:
                if i > 13: currbrigade.append((i-1,brigades[i-4]))
                else: currbrigade.append((i-1,brigades[i-2]))
    elif clm == 0:
        for i in range(6):
            currbrigade.append((0,'#NAME SURNAME\n#NAME SURNAME'))

    #TITLE CREATION PART 3

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1r1 = p1.add_run('Министерство науки и высшего образования\n Российской Федерации Мытищинский филиал\n')
    path_to_picture = working_directory / "media/images/herb.jpg"
    # print(str(path_to_picture))
    p1r1.add_picture(str(path_to_picture))
    p1r1.add_text('\nФедеральное государственное бюджетное образовательное\n учреждение\n высшего образования\n' 
                   + '«Московский государственный технический университет\n имени Н.Э. Баумана\n' 
                   + '(национальный исследовательский университет)»\n(МГТУ им. Н.Э. Баумана)\n')
    p1r1.add_text('______________________________________________________________________________________________\n'
                  + '______________________________________________________________________________________________\n')
    p1r1.font.bold=True
    p1r1.font.size = Pt(12)
    
    p1r2 = p1.add_run('ФАКУЛЬТЕТ Космический\n КАФЕДРА «Прикладная математика, информатика и вычислительная техника» K3-МФ\n\n\n')
    p1r2.font.size = Pt(14)
    p1r2.font.bold = True

    p1r3 = p1.add_run('Отчет\n По дежурству\n\n')
    p1r3.font.size = Pt(28)
    p1r3.font.bold = True

    p2r1 = p2.add_run('Бригада: '+ str(currbrigade[0][0])+'\n'+currbrigade[0][1]+'\n')
    p2r1.add_text('Бригада: '+str(currbrigade[1][0])+'\n'+currbrigade[1][1]+'\n')
    p2r1.add_text('Бригада: '+str(currbrigade[2][0])+ '\n'+currbrigade[2][1]+'\n\n')
    p2r1.font.size = Pt(14)

    p3r1 = p3.add_run('Дата прохождения дежурства: '+current_date[0:10]+'\n\n\n'+'Москва '+current_date[0:4]+'\n\n')
    p3r1.font.size = Pt(14)
    p3r1.font.bold = True
    p3r1.add_break(WD_BREAK.PAGE)

    #TABLE CREATION PART 4

    tablecount = (len(dates)-1)/5
    if tablecount%1!=0: tablecount = int(tablecount)+1

    columns = 6
    if len(dates)<=5: columns=len(dates)

    for i in range(int(tablecount)):
        tables.append(doc.add_table(df.shape[1] + 1, columns))
        tables[i].cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        trun = tables[i].cell(0,0).paragraphs[0].add_run('Дата')
        trun.font.bold=True
        brun = doc.add_paragraph().add_run()
        brun.add_break(docx.enum.text.WD_BREAK.PAGE)

    k = 0
    m = 0
    for i in range(len(tables)):
        for j in range(df.shape[1]):
            for l in range(int(tablecount)):shade_col.append(parse_xml(r'<w:shd {} w:fill="E6E6E6"/>'.format(nsdecls('w'))))
            tables[i].cell(j+1,0)._tc.get_or_add_tcPr().append(shade_col[m])
            paragraph = tables[i].cell(j + 1, 0).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            trun = paragraph.add_run(df.columns[j])
            trun.font.bold = True
            trun.font.size = Pt(8)
            m=m+1

        for j in range(columns-1):
            if(k+1>len(dates)-1):break
            for l in range(int(tablecount)):shade_row.append(parse_xml(r'<w:shd {} w:fill="E6E6E6"/>'.format(nsdecls('w'))))
            tables[i].cell(0,j+1)._tc.get_or_add_tcPr().append(shade_row[k])
            paragraph = tables[i].cell(0, j + 1).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            trun = paragraph.add_run(dates[k].strftime("%H:%M:%S - ") + dates[k+1].strftime("%H:%M:%S"))
            trun.font.bold = True
            trun.font.size = Pt(8)
            k=k+1
    
    k = 0
    for i, list_of_vals in enumerate(values):
        k=int(i/(columns-1))
        for j, val in enumerate(list_of_vals):
            if val == 0:
                paragraph = tables[k].cell(j+1, (i%(columns-1)) + 1).paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run('✖')
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.bold = True

            elif val==1:
                paragraph = tables[k].cell(j+1, (i%(columns-1)) + 1).paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run('✔')
                run.font.color.rgb = RGBColor(0, 200, 0)
                run.font.bold = True

    doc.save(path_file_out)

if __name__ == '__main__':
    main()
